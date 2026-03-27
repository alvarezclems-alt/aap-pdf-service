"""
api.py — Service Railway complet DocAdmin
PDF AAP + PDF Vacataire + IA Claude + Remplissage + Révision de documents
"""
import io, os, json, traceback, base64, re
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS


app = Flask(__name__)
CORS(app, origins="*")

# ── Imports optionnels ─────────────────────────────────────────────────────────
try:
    import anthropic
    claude = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))
    CLAUDE_MODEL = "claude-sonnet-4-20250514"
    CLAUDE_OK = True
except ImportError:
    CLAUDE_OK = False
    print("WARNING: anthropic non installé")

try:
    import fitz  # PyMuPDF
    FITZ_OK = True
except ImportError:
    FITZ_OK = False
    print("WARNING: PyMuPDF (fitz) non installé — remplissage PDF via LibreOffice")

try:
    from generate_annexes import build_annexe1, build_annexe1bis
    def generate_annexe1_pdf(data):
        import os
        logo_path = os.path.join(os.path.dirname(__file__), 'logo_inspe.png')
        logo_bytes = None
        if os.path.exists(logo_path):
            with open(logo_path, 'rb') as f:
                logo_bytes = f.read()
        project = data.get('project', data)
        return build_annexe1(project, logo_bytes)

    def generate_annexe1bis_pdf(data):
        import os
        logo_path = os.path.join(os.path.dirname(__file__), 'logo_inspe.png')
        logo_bytes = None
        if os.path.exists(logo_path):
            with open(logo_path, 'rb') as f:
                logo_bytes = f.read()
        project = data.get('project', data)
        budget_lines = data.get('budget_lines', data.get('lignes', []))
        return build_annexe1bis(project, budget_lines, logo_bytes)

    ANNEXES_OK = True
except ImportError as e:
    ANNEXES_OK = False
    print(f"WARNING: generate_annexes non trouvé: {e}")

try:
    from generate_vacataire import generate_vacataire_pdf
    VACATAIRE_OK = True
except ImportError as e:
    VACATAIRE_OK = False
    print(f"WARNING: generate_vacataire non trouvé: {e}")

try:
    from docx import Document as DocxDoc
    DOCX_OK = True
except ImportError as e:
    DOCX_OK = False
    print(f"WARNING: python-docx non trouvé: {e}")


# ── Utilitaire Claude ──────────────────────────────────────────────────────────
def claude_text(system, user, max_tokens=800):
    if not CLAUDE_OK:
        return "Service IA non disponible"
    import time
    last_error = None
    for tentative in range(3):
        try:
            msg = claude.messages.create(
                model=CLAUDE_MODEL, max_tokens=max_tokens,
                system=system,
                messages=[{"role": "user", "content": user}]
            )
            return msg.content[0].text.strip()
        except Exception as e:
            last_error = e
            if "529" in str(e) or "overloaded" in str(e).lower():
                print(f"[CLAUDE] API surchargée, tentative {tentative+1}/3, attente {2**tentative}s")
                time.sleep(2 ** tentative)
            else:
                raise
    raise last_error


# ── Health ─────────────────────────────────────────────────────────────────────
@app.route("/health", methods=["GET"])
def health():
    return jsonify({
        "status": "ok",
        "modules": {
            "claude":    CLAUDE_OK,
            "annexes":   ANNEXES_OK,
            "vacataire": VACATAIRE_OK,
            "docx":      DOCX_OK,
            "fitz":      FITZ_OK,
        },
        "anthropic_key_set": bool(os.environ.get("ANTHROPIC_API_KEY")),
        "endpoints": [
            "POST /generate-annexe1",
            "POST /generate-annexe1bis",
            "POST /generate-all",
            "POST /generate-vacataire",
            "POST /ai/justification-vacataire",
            "POST /ai/generer-section-aap",
            "POST /ai/reviser-section",
            "POST /ai/reviser-document",
            "POST /ai/extraire-infos",
            "POST /ai/remplir-document",
        ]
    }), 200


# ══════════════════════════════════════════════════════════════════════════════
# PDF AAP
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/generate-annexe1", methods=["POST"])
def generate_annexe1():
    if not ANNEXES_OK:
        return jsonify({"error": "Module generate_annexes non disponible"}), 503
    try:
        data = request.get_json(force=True)
        pdf  = generate_annexe1_pdf(data)
        return send_file(io.BytesIO(pdf), mimetype="application/pdf",
                         as_attachment=True, download_name="annexe1.pdf")
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route("/generate-annexe1bis", methods=["POST"])
def generate_annexe1bis():
    if not ANNEXES_OK:
        return jsonify({"error": "Module generate_annexes non disponible"}), 503
    try:
        data = request.get_json(force=True)
        pdf  = generate_annexe1bis_pdf(data)
        return send_file(io.BytesIO(pdf), mimetype="application/pdf",
                         as_attachment=True, download_name="annexe1bis.pdf")
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route("/generate-all", methods=["POST"])
def generate_all():
    import zipfile
    if not ANNEXES_OK:
        return jsonify({"error": "Module generate_annexes non disponible"}), 503
    try:
        data = request.get_json(force=True)
        buf  = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
            z.writestr("annexe1.pdf",    generate_annexe1_pdf(data))
            z.writestr("annexe1bis.pdf", generate_annexe1bis_pdf(data))
        buf.seek(0)
        return send_file(buf, mimetype="application/zip",
                         as_attachment=True, download_name="dossier-complet.zip")
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ══════════════════════════════════════════════════════════════════════════════
# PDF VACATAIRE
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/generate-vacataire", methods=["POST"])
def generate_vacataire():
    if not VACATAIRE_OK:
        return jsonify({"error": "Module generate_vacataire non disponible"}), 503
    try:
        data = request.get_json(force=True) or {}
        i    = data.get("intervenant", {})
        if not i.get("nom") and not i.get("prenom"):
            return jsonify({"error": "nom/prénom manquant"}), 400
        pdf    = generate_vacataire_pdf(data)
        nom    = i.get("nom", "intervenant").replace(" ", "-").lower()
        prenom = i.get("prenom", "").replace(" ", "-").lower()
        return send_file(io.BytesIO(pdf), mimetype="application/pdf",
                         as_attachment=True,
                         download_name=f"dossier-vacataire-{prenom}-{nom}.pdf")
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ══════════════════════════════════════════════════════════════════════════════
# IA — JUSTIFICATION VACATAIRE
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/ai/justification-vacataire", methods=["POST"])
def ai_justification_vacataire():
    if not CLAUDE_OK:
        return jsonify({"error": "Service Claude non disponible"}), 503
    try:
        d = request.get_json(force=True) or {}
        system = (
            "Tu es un assistant administratif expert dans la rédaction "
            "de dossiers universitaires français. Ton style est factuel, "
            "professionnel et précis. Tu rédiges uniquement en français."
        )
        user = f"""Rédige en 6 à 8 lignes un texte justifiant la compétence
d'un intervenant vacataire. Commence directement par les compétences,
sans phrase d'introduction. Utilise la troisième personne.

Intervenant : {d.get('prenom','')} {d.get('nom','')}
Profil : {d.get('type_profil','')}
Spécialités : {d.get('specialites','')}
Situation : {d.get('situation_pro','')}
Employeur : {d.get('employeur_nom','') or 'Indépendant'}

Mission :
- Intitulé : {d.get('intitule','')}
- Nature : {d.get('nature_intervention','')}
- Niveau : {d.get('niveau','')}
- Établissement : {d.get('etablissement_nom','')}

Texte :"""
        return jsonify({"texte": claude_text(system, user, 600)})
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ══════════════════════════════════════════════════════════════════════════════
# IA — GÉNÉRER SECTION AAP
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/ai/generer-section-aap", methods=["POST"])
def ai_generer_section_aap():
    if not CLAUDE_OK:
        return jsonify({"error": "Service Claude non disponible"}), 503
    try:
        d       = request.get_json(force=True) or {}
        section = d.get("section", "description")
        p       = d.get("projet", {})
        system  = (
            "Tu es un expert en rédaction de dossiers de candidature "
            "pour des appels à projets de recherche en éducation français. "
            "Ton style est académique, structuré et convaincant. "
            "Tu rédiges uniquement en français."
        )
        prompts = {
            "resume": f"""Rédige un résumé de 8 à 10 lignes pour ce projet.
Titre : {p.get('titre','')}
Description : {p.get('description_libre','')}
Axe : {p.get('axe','')} | Mots-clefs : {p.get('mots_clefs','')}
Présente : contexte, approche, résultats attendus, livrables.
Résumé :""",

            "description": f"""Rédige la description scientifique complète (3 pages max).
Structure avec ces sous-titres en gras :
**Inscription générale et problématique**
**Résultats antérieurs, originalité et enjeux**
**Objectifs précis et hypothèses de travail**
**Cadres théoriques, méthodologie et sources**
**Résultats attendus, restitution et diffusion**

Titre : {p.get('titre','')}
Description : {p.get('description_libre','')}
Axe : {p.get('axe','')} | Budget : {p.get('montant','')} €
Description :""",

            "calendrier": f"""Génère un calendrier en tableau Markdown pour ce projet
de recherche sur 2 ans (2027-2028). Format obligatoire :
| Grandes étapes | Début prévisionnel | Fin prévisionnelle | Durée estimée |
|---|---|---|---|
Titre : {p.get('titre','')}
8 à 10 étapes couvrant : littérature, collecte, analyse, restitution, livrables, valorisation.
Tableau :""",

            "budget": f"""Justifie les dépenses pour un budget de {p.get('montant',8000)} €
sur 2 ans pour ce projet : {p.get('titre','')}
Répartition : missions, prestations, séminaires, vacations, matériel.
Format : liste avec montants en euros.
Justification :""",
        }
        max_t = 2000 if section == "description" else 800
        texte = claude_text(system, prompts.get(section, prompts["description"]), max_t)
        return jsonify({"texte": texte, "section": section})
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ══════════════════════════════════════════════════════════════════════════════
# IA — RÉVISER UNE SECTION (champ texte)
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/ai/reviser-section", methods=["POST"])
def ai_reviser_section():
    if not CLAUDE_OK:
        return jsonify({"error": "Service Claude non disponible"}), 503
    try:
        d = request.get_json(force=True) or {}
        system = (
            "Tu es un expert en rédaction académique française. "
            "Tu révises les textes en respectant les instructions. "
            "Tu retournes uniquement le texte révisé, sans commentaire."
        )
        user = f"""Contexte : {d.get('contexte', 'Dossier administratif')}
Instruction : {d.get('instruction', 'Améliore ce texte')}
Texte original :
{d.get('texte_original', '')}
Texte révisé :"""
        return jsonify({"texte": claude_text(system, user, 1500)})
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ══════════════════════════════════════════════════════════════════════════════
# IA — RÉVISER UN DOCUMENT AAP COMPLET (chat)
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/ai/reviser-document", methods=["POST"])
def ai_reviser_document():
    if not CLAUDE_OK:
        return jsonify({"error": "Service Claude non disponible"}), 503
    try:
        d           = request.get_json(force=True) or {}
        texte       = d.get("texte_document", "")
        instruction = d.get("instruction", "")
        historique  = d.get("historique", [])
        projet      = d.get("projet", {})

        if not texte:
            return jsonify({"error": "texte_document manquant"}), 400
        if not instruction:
            return jsonify({"error": "instruction manquante"}), 400

        system = (
            "Tu es un assistant expert en rédaction de dossiers de candidature "
            "pour des appels à projets de recherche en éducation français. "
            "Quand on te demande de modifier un document, tu appliques les "
            "changements demandés avec précision et tu retournes le texte "
            "complet du document modifié. "
            "Tu rédiges uniquement en français. "
            "Tu retournes UNIQUEMENT un objet JSON valide, "
            "sans texte avant ni après, sans balises markdown."
        )

        messages = []
        for msg in historique[-6:]:
            role    = msg.get("role", "user")
            content = msg.get("content", "")
            if role in ("user", "assistant") and content:
                messages.append({"role": role, "content": content})

        messages.append({
            "role": "user",
            "content": f"""Voici le document AAP actuel :

---
{texte[:8000]}
---

Projet : {projet.get('titre', '')}
Description : {projet.get('description', '')[:500]}

Instruction de modification : {instruction}

Applique cette modification et retourne un JSON avec exactement ces deux clés :
{{
  "reponse": "Résumé clair en 2-3 lignes des modifications apportées",
  "texte_modifie": "Texte COMPLET du document après modification (conserve toute la structure)"
}}

JSON :"""
        })

        response = claude.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=4000,
            system=system,
            messages=messages
        )

        raw = response.content[0].text.strip()
        if raw.startswith("```"):
            parts = raw.split("```")
            raw = parts[1] if len(parts) > 1 else raw
            if raw.startswith("json"):
                raw = raw[4:]
        raw = raw.strip()

        try:
            result = json.loads(raw)
            if "reponse" not in result:
                result["reponse"] = "Document modifié avec succès."
            if "texte_modifie" not in result:
                result["texte_modifie"] = texte
            return jsonify(result)
        except json.JSONDecodeError:
            return jsonify({
                "reponse": raw[:500] if raw else "Modification appliquée.",
                "texte_modifie": texte
            })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ══════════════════════════════════════════════════════════════════════════════
# IA — EXTRAIRE INFOS D'UN TEXTE
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/ai/extraire-infos", methods=["POST"])
def ai_extraire_infos():
    if not CLAUDE_OK:
        return jsonify({"error": "Service Claude non disponible"}), 503
    raw = ""
    try:
        d     = request.get_json(force=True) or {}
        texte = d.get("texte", "")
        typ   = d.get("type", "projet")
        system = (
            "Tu es un assistant expert en extraction d'informations. "
            "Tu retournes UNIQUEMENT un objet JSON valide, "
            "sans texte avant ni après, sans balises markdown."
        )
        schema = (
            '{"titre":"","acronyme":"","mots_clefs":"","resume":"",'
            '"description":"","coordinateur_nom":"","coordinateur_email":"",'
            '"institution":"","laboratoire":"","membres":[],'
            '"montant":0,"axe":""}') if typ == "projet" else (
            '{"nom_aap":"","institution":"","deadline":"",'
            '"budget_max":0,"axes":[],"criteres_eligibilite":"","description":""}')
        user = f"Extrais les informations et retourne ce JSON :\n{schema}\n\nTexte :\n{texte[:6000]}\n\nJSON :"
        raw  = claude_text(system, user, 1200)
        raw  = raw.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        return jsonify({"infos": json.loads(raw.strip())})
    except json.JSONDecodeError as e:
        return jsonify({"error": f"JSON invalide: {e}", "raw": raw}), 500
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ══════════════════════════════════════════════════════════════════════════════
# IA — EXTRAIRE JUSTIFICATIF
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/ai/extraire-justificatif", methods=["POST"])
def ai_extraire_justificatif():
    if not CLAUDE_OK:
        return jsonify({"error": "Service Claude non disponible"}), 503
    try:
        d = request.get_json(force=True) or {}
        b64 = d.get("fichier_base64", "")
        mime = d.get("type_mime", "")
        nom = d.get("nom_fichier", "fichier")

        if not b64:
            return jsonify({"error": "Fichier manquant"}), 400

        system = (
            "Tu es un expert comptable et administratif français. "
            "Tu analyses des justificatifs de voyage (billets SNCF, factures, RIB, etc.) "
            "et tu retournes UNIQUEMENT un JSON valide, sans texte avant ni après, "
            "sans balises markdown."
        )
        user = f"""Analyse ce document ({nom}) qui peut contenir UN OU PLUSIEURS billets/trajets.

RÈGLES IMPORTANTES :
- Si le document contient PLUSIEURS billets ou trajets (ex: billet aller + billet retour, ou plusieurs pages),
  additionne TOUS les montants et retourne le TOTAL dans "montant"
- Pour les billets SNCF : type_document = "billet_train" (même pour TGV, Intercités, etc.)
- La date doit être au format JJ/MM/AA (ex: 10/03/26)
- Le montant doit être un nombre décimal sans symbole (ex: 112.00 et non "112,00 €")

Retourne ce JSON (ne mets que les champs que tu trouves) :
{{
  "type_document": "billet_train" | "billet_avion" | "facture_hotel" | "taxi" | "peage" | "rib" | "autre",
  "resume": "Description courte (ex: Billets SNCF Tours-Nantes-Paris-Lille)",
  "informations": {{
    "montant": NOMBRE_TOTAL_FLOTTANT,
    "devise": "EUR",
    "date": "JJ/MM/AA",
    "origine": "Ville de départ du premier trajet",
    "destination": "Ville d'arrivée du dernier trajet",
    "nb_billets": NOMBRE_DE_BILLETS_DANS_LE_DOCUMENT,
    "detail_montants": "ex: 36€ + 37€ = 73€",
    "iban": "FR...",
    "bic": "...",
    "titulaire": "Nom du titulaire"
  }}
}}"""

        content = []
        if mime == "application/pdf":
            content.append({"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": b64}})
        elif mime in ["image/jpeg", "image/png", "image/webp", "image/gif"]:
            content.append({"type": "image", "source": {"type": "base64", "media_type": mime, "data": b64}})
        elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or nom.endswith(".docx"):
            docx_bytes = base64.b64decode(b64)
            doc = DocxDoc(io.BytesIO(docx_bytes))
            texte = "\n".join([p.text for p in doc.paragraphs])
            content.append({"type": "text", "text": f"Contenu du DOCX :\n{texte[:10000]}\n\n"})
        else:
            return jsonify({"error": f"Format non supporté par l'IA: {mime}"}), 400

        content.append({"type": "text", "text": user})

        import time
        last_error = None
        for tentative in range(3):  # 3 tentatives max
            try:
                response = claude.messages.create(
                    model=CLAUDE_MODEL, max_tokens=1000, system=system,
                    messages=[{"role": "user", "content": content}]
                )
                break  # Succès → sortir de la boucle
            except Exception as e:
                last_error = e
                if "529" in str(e) or "overloaded" in str(e).lower():
                    print(f"[JUSTIF] API surchargée, tentative {tentative+1}/3, attente {2**tentative}s")
                    time.sleep(2 ** tentative)  # 1s, 2s, 4s
                else:
                    raise  # Autre erreur → remonter immédiatement
        else:
            raise last_error  # 3 échecs → erreur

        raw = response.content[0].text.strip()
        if raw.startswith("```"):
            parts = raw.split("```")
            raw = parts[1] if len(parts) > 1 else raw
            if raw.startswith("json"):
                raw = raw[4:]
        raw = raw.strip()
        return jsonify(json.loads(raw))
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ══════════════════════════════════════════════════════════════════════════════
# UTILITAIRES — CONVERSION FICHIERS
# ══════════════════════════════════════════════════════════════════════════════

def _decomposer_adresse(adresse: str):
    adresse = adresse.strip()
    m = re.match(r'^(\d+\s*(?:bis|ter|quater)?)\s+(.+)$', adresse, re.IGNORECASE)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return "", adresse


def _fusionner_runs(para):
    if len(para.runs) <= 1:
        return
    texte_complet = para.text
    if not texte_complet.strip():
        return
    para.runs[0].text = texte_complet
    for run in para.runs[1:]:
        run.text = ""


def _convert_pdf_to_docx(pdf_bytes: bytes, nom_fich: str) -> bytes:
    import subprocess, tempfile, glob
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = os.path.join(tmpdir, nom_fich)
        with open(input_path, "wb") as f:
            f.write(pdf_bytes)
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", tmpdir, input_path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice erreur: {result.stderr}")
        docx_files = glob.glob(os.path.join(tmpdir, "*.docx"))
        if not docx_files:
            raise RuntimeError("Conversion PDF→DOCX : aucun fichier produit")
        with open(docx_files[0], "rb") as f:
            return f.read()


def _convert_doc_to_docx(doc_bytes: bytes, nom_fich: str) -> bytes:
    import subprocess, tempfile, glob
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = os.path.join(tmpdir, nom_fich)
        with open(input_path, "wb") as f:
            f.write(doc_bytes)
        ext = os.path.splitext(nom_fich)[1].lower()
        cmd = ["libreoffice", "--headless"]
        if ext == ".pdf":
            cmd.extend(["--infilter=writer_pdf_import"])
        cmd.extend(["--convert-to", "docx", "--outdir", tmpdir, input_path])
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice erreur: {result.stderr}")
        docx_files = glob.glob(os.path.join(tmpdir, "*.docx"))
        if not docx_files:
            raise RuntimeError("Conversion LibreOffice : aucun .docx produit")
        with open(docx_files[0], "rb") as f:
            return f.read()


def _convert_docx_to_pdf(docx_bytes: bytes, nom_fich: str) -> bytes:
    import subprocess, tempfile, glob
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = os.path.join(tmpdir, "doc.docx")
        with open(input_path, "wb") as f:
            f.write(docx_bytes)
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, input_path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice erreur: {result.stderr}")
        pdf_files = glob.glob(os.path.join(tmpdir, "*.pdf"))
        if not pdf_files:
            raise RuntimeError("Conversion LibreOffice : aucun .pdf produit")
        with open(pdf_files[0], "rb") as f:
            return f.read()


def _construire_justif_bloc(justifs: list) -> str:
    """
    Construit un bloc texte clair pour les prompts Claude.
    Inclut les totaux consolidés par catégorie ET le détail de chaque justificatif.
    """
    if not justifs:
        return ""

    # Consolider par catégorie
    totaux = {}
    dates = []
    rib_info = {}

    for j in justifs:
        ttype = j.get("type_document", "autre")
        info = j.get("informations", {})
        montant = info.get("montant")
        date = info.get("date", "")

        if montant is not None:
            try:
                m = float(montant)
                totaux[ttype] = totaux.get(ttype, 0.0) + m
            except (ValueError, TypeError):
                pass

        if date:
            dates.append(date)

        if ttype == "rib":
            rib_info = info

    # Construire le résumé consolidé
    lignes_consolidees = []
    if "billet_train" in totaux:
        lignes_consolidees.append(f"  → TOTAL TRAIN = {totaux['billet_train']:.2f} EUR")
    if "billet_avion" in totaux:
        lignes_consolidees.append(f"  → TOTAL AVION = {totaux['billet_avion']:.2f} EUR")
    if "taxi" in totaux:
        lignes_consolidees.append(f"  → TOTAL TAXI = {totaux['taxi']:.2f} EUR")
    if "peage" in totaux:
        lignes_consolidees.append(f"  → TOTAL PÉAGE = {totaux['peage']:.2f} EUR")

    total_global = sum(totaux.values())
    if len(totaux) > 1:
        lignes_consolidees.append(f"  → MONTANT TOTAL GLOBAL = {total_global:.2f} EUR")

    if dates:
        lignes_consolidees.append(f"  → DATE DÉPART = {dates[0]}")
        if len(dates) > 1:
            lignes_consolidees.append(f"  → DATE ARRIVÉE = {dates[-1]}")

    if rib_info.get("iban"):
        lignes_consolidees.append(f"  → IBAN = {rib_info['iban']}")
    if rib_info.get("bic"):
        lignes_consolidees.append(f"  → BIC = {rib_info['bic']}")

    # Détail de chaque justificatif
    detail_lines = []
    for j in justifs:
        info = j.get("informations", {})
        detail_lines.append(
            f"  - [{j.get('type_document','')}] {j.get('resume','')} "
            f"montant={info.get('montant','')} date={info.get('date','')} "
            f"origine={info.get('origine','')} destination={info.get('destination','')}"
        )

    bloc = "\nJustificatifs fournis (UTILISE CES DONNÉES POUR REMPLIR LE FORMULAIRE) :\n"
    if lignes_consolidees:
        bloc += "Totaux calculés :\n" + "\n".join(lignes_consolidees) + "\n"
    bloc += "Détail :\n" + "\n".join(detail_lines) + "\n"
    return bloc


# ══════════════════════════════════════════════════════════════════════════════
# IA — REMPLIR UN DOCUMENT ADMINISTRATIF (PDF, DOC, DOCX)
# ══════════════════════════════════════════════════════════════════════════════

def _inserer_valeur_pdf(page, x, y, valeur, fontsize=9):
    """Insère une valeur à une position (x, y) sur la page PDF."""
    if not valeur:
        return
    page.insert_text(
        (x, y),
        str(valeur),
        fontsize=fontsize,
        color=(0, 0, 0),
        fontname="helv"
    )


def _effacer_zone(page, x0, y0, x1, y1):
    """Efface une zone rectangulaire (rectangle blanc)."""
    import fitz as _fitz
    rect = _fitz.Rect(x0, y0, x1, y1)
    page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))


def _remplir_pdf_par_coordonnees(page, profil, marqueurs_coords):
    """
    Remplit les zones connues du PDF par coordonnées fixes.
    marqueurs_coords = liste de (marqueur_texte, x_insertion, y_insertion, valeur)
    """
    for marqueur, x_ins, y_ins, valeur in marqueurs_coords:
        if not valeur:
            continue
        # Effacer le marqueur existant
        instances = page.search_for(marqueur)
        for inst in instances:
            page.draw_rect(inst, color=(1, 1, 1), fill=(1, 1, 1))
        # Insérer la valeur
        _inserer_valeur_pdf(page, x_ins, y_ins, valeur)


def _detecter_zones_pdf(pdf_bytes):
    """
    Détecte dynamiquement toutes les zones à remplir dans un PDF.
    Retourne une liste de zones avec leur label détecté et leurs coordonnées.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    zones = []

    MARQUEURS_VIDES = [
        "…………………………………… Euros",
        "………………………………………",
        "…………………………………",
        "………………………………",
        "……………………………",
        "…………………………",
        "………………………",
        "……………………",
        "…………………",
        "………………",
        "……………",
        "…………",
        "………",
        "……",
        "............................................................................",
        "...................................................................",
        "................................................................",
        ".............................................................",
        ".........................................................",
        "......................................................",
        "...................................................",
        "................................................",
        ".............................................",
        "...........................................",
        ".........................................",
        ".......................................",
        ".....................................",
        "...................................",
        ".................................",
        "...............................",
        ".............................",
        "...........................",
        ".........................",
        ".......................",
        ".....................",
        "...................",
        ".................",
        "...............",
        ".............",
        "...........",
        ".........",
        ".......",
        "......",
        "(jj/mm/aa)",
        "(jj/mm/aaaa)",
        "____________________",
        "__________________",
        "________________",
        "______________",
    ]

    for page_num, page in enumerate(doc):
        # Extraire tous les mots avec leurs positions
        words = page.get_text("words")  # (x0, y0, x1, y1, word, block, line, word_idx)

        # Chercher chaque marqueur
        zones_trouvees_y = set()  # éviter doublons sur même ligne

        for marqueur in MARQUEURS_VIDES:
            instances = page.search_for(marqueur)
            for inst in instances:
                y_key = round(inst.y0)
                if y_key in zones_trouvees_y:
                    continue
                zones_trouvees_y.add(y_key)

                # Trouver le label le plus proche à gauche ou au-dessus
                label = _trouver_label_proche(words, inst, page)

                zones.append({
                    "page": page_num,
                    "marqueur": marqueur,
                    "rect": [inst.x0, inst.y0, inst.x1, inst.y1],
                    "label": label,
                    "x_insert": inst.x0 + 1,
                    "y_insert": inst.y1 - 1,
                })

    doc.close()
    return zones


def _trouver_label_proche(words, rect, page):
    """
    Trouve le label textuel le plus proche à gauche ou au-dessus d'une zone.
    Exclut les marqueurs de zones vides.
    """
    MOTS_EXCLUS = {
        "euros", "eur", "€", "x", "0,5", "km", ":",
        "………", "......", "____", "(jj/mm/aa)", "(jj/mm/aaaa)"
    }

    # Chercher les mots sur la même ligne (±8px) à gauche
    candidats_gauche = []
    for w in words:
        wx0, wy0, wx1, wy1, word = w[0], w[1], w[2], w[3], w[4]
        mot = word.strip().lower().rstrip(":").strip()
        if mot in MOTS_EXCLUS or len(mot) < 2:
            continue
        if abs(wy0 - rect.y0) < 10 and wx1 <= rect.x0 + 5:
            dist = rect.x0 - wx1
            candidats_gauche.append((dist, word.strip()))

    if candidats_gauche:
        # Prendre le mot le plus proche à gauche
        candidats_gauche.sort(key=lambda x: x[0])
        # Reconstruire le label depuis les mots proches
        mots_label = []
        for dist, mot in candidats_gauche[:4]:
            if dist < 150:
                mots_label.insert(0, mot)
        return " ".join(mots_label).strip().rstrip(":")

    # Sinon chercher au-dessus (±30px)
    candidats_dessus = []
    for w in words:
        wx0, wy0, wx1, wy1, word = w[0], w[1], w[2], w[3], w[4]
        mot = word.strip().lower().rstrip(":").strip()
        if mot in MOTS_EXCLUS or len(mot) < 2:
            continue
        if wy1 <= rect.y0 and wy1 >= rect.y0 - 30:
            dist = rect.y0 - wy1
            candidats_dessus.append((dist, word.strip()))

    if candidats_dessus:
        candidats_dessus.sort(key=lambda x: x[0])
        return candidats_dessus[0][1].rstrip(":")

    return ""


def _remplir_pdf_direct(pdf_bytes, nom_fich, profil_enrichi, justifs):
    """
    Remplit un PDF dynamiquement — fonctionne avec N'IMPORTE QUEL formulaire PDF.

    Stratégie :
    1. Détecter automatiquement toutes les zones à remplir (marqueurs + coordonnées)
    2. Envoyer à Claude : zones détectées + profil + justificatifs
    3. Claude retourne le mapping zone → valeur
    4. Insérer les valeurs aux coordonnées exactes avec PyMuPDF
    """
    p = profil_enrichi

    # ── 1. Détecter toutes les zones à remplir ────────────────────────────────
    zones = _detecter_zones_pdf(pdf_bytes)

    if not zones:
        print("[REMPLIR_PDF] Aucune zone détectée, retour PDF original")
        return pdf_bytes

    # ── 2. Construire la représentation pour Claude ───────────────────────────
    zones_str = "\n".join(
        f"[{i}] label='{z['label']}' marqueur='{z['marqueur'][:20]}' "
        f"x={z['x_insert']:.0f} y={z['y_insert']:.0f} page={z['page']}"
        for i, z in enumerate(zones)
    )

    justif_bloc = _construire_justif_bloc(justifs)

    system = (
        "Tu es un assistant expert en administration française. "
        "Tu analyses des formulaires et associes les zones vides aux bonnes valeurs du profil. "
        "Tu retournes UNIQUEMENT un objet JSON valide, sans texte avant ni après, sans balises markdown."
    )

    user = f"""Voici les zones vides détectées dans un formulaire PDF :
(chaque zone a un label détecté automatiquement et des coordonnées d'insertion)

{zones_str}

Profil de la personne :
{json.dumps(p, ensure_ascii=False, indent=2)}
{justif_bloc}

Retourne un JSON avec les valeurs à insérer pour chaque zone :
{{
  "remplissages": [
    {{"index": 0, "valeur": "DUPONT"}},
    {{"index": 1, "valeur": "Marie"}}
  ]
}}

RÈGLES ABSOLUES :
1. "index" = index [N] de la zone dans la liste ci-dessus
2. "valeur" = UNIQUEMENT la donnée à insérer

MAPPING NOM/PRÉNOM :
- label contient "nom" (sans "prénom") → profil["nom"] = NOM DE FAMILLE uniquement
- label contient "prénom" → profil["prenom"] = PRÉNOM uniquement
- Ne jamais mettre nom+prénom ensemble dans un seul champ

MAPPING TABLEAU VOYAGE (labels "Train", "Avion", "Taxi", etc.) :
- label "Train" ou contient "train" → profil["frais_train"] + " Euros"
- label "Avion" → profil["frais_avion"] + " Euros"
- label "Taxi" → profil["frais_taxi"] + " Euros"
- label "Péage" → profil["frais_peage"] + " Euros"
- label "Voiture" ou "km" → profil["frais_voiture"] + " Euros"
- label "Montant total" ou "total" → profil["montant_total"] + " Euros"
- label "Départ" ou "départ" → profil["date_depart"]
- label "Arrivée" ou "arrivée" → profil["date_arrivee"]
- label "Autre" → laisser vide sauf si frais "autre" existent

AUTRES MAPPINGS COURANTS :
- label "ville" → profil["ville"]
- label "téléphone" ou "tél" → profil["telephone"]
- label "email" ou "e-mail" ou "mail" → profil["email"]
- label "civilité" ou cases □ Pr □ Dr □ M → profil["civilite"]
- label "iban" → profil["iban"]
- label "date" (champ signature) → date du jour

Si une valeur n'est pas dans le profil → ne pas inclure cet index.

JSON :"""

    import time
    last_error = None
    raw = ""
    for tentative in range(3):
        try:
            response = claude.messages.create(
                model=CLAUDE_MODEL, max_tokens=2000,
                system=system,
                messages=[{"role": "user", "content": user}]
            )
            raw = response.content[0].text.strip()
            break
        except Exception as e:
            last_error = e
            if "529" in str(e) or "overloaded" in str(e).lower():
                print(f"[REMPLIR_PDF] API surchargée tentative {tentative+1}/3")
                time.sleep(2 ** tentative)
            else:
                raise
    else:
        raise last_error

    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1] if len(parts) > 1 else raw
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip()

    try:
        result = json.loads(raw)
        remplissages = {str(r["index"]): r["valeur"] for r in result.get("remplissages", [])}
    except (json.JSONDecodeError, KeyError):
        print(f"[REMPLIR_PDF] JSON invalide: {raw[:200]}")
        remplissages = {}

    print(f"[REMPLIR_PDF] {len(zones)} zones détectées, {len(remplissages)} remplissages")

    # ── 3. Gérer la civilité séparément (cases à cocher) ─────────────────────
    civilite_valeur = p.get("civilite", "")

    # ── 4. Appliquer les insertions avec PyMuPDF ──────────────────────────────
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    for idx_str, valeur in remplissages.items():
        try:
            idx = int(idx_str)
            if idx >= len(zones):
                continue
            zone = zones[idx]
            valeur = str(valeur).strip()
            if not valeur:
                continue

            page = doc[zone["page"]]
            marqueur = zone["marqueur"]
            rect_zone = zone["rect"]

            # Effacer le marqueur
            instances = page.search_for(marqueur)
            for inst in instances:
                if abs(inst.y0 - rect_zone[1]) < 10:
                    page.draw_rect(inst, color=(1, 1, 1), fill=(1, 1, 1))
                    break

            # Insérer la valeur
            fontsize = max(7, min(10, (rect_zone[3] - rect_zone[1]) * 0.75))
            page.insert_text(
                (zone["x_insert"], zone["y_insert"]),
                valeur,
                fontsize=fontsize,
                color=(0, 0, 0),
                fontname="helv"
            )

        except Exception as e:
            print(f"[REMPLIR_PDF] Erreur zone {idx_str}: {e}")

    # ── 5. Gérer les cases à cocher civilité ─────────────────────────────────
    if civilite_valeur:
        for page_num in range(len(doc)):
            page = doc[page_num]
            CIVILITES_MAP = {
                "pr": "q Pr", "dr": "q Dr",
                "m.": "q M.", "m": "q M.",
                "mme": "q Mme", "mlle": "q Mlle"
            }
            civ_key = civilite_valeur.lower().strip().rstrip(".")
            marqueur_civ = CIVILITES_MAP.get(civ_key + ".") or CIVILITES_MAP.get(civ_key)
            if marqueur_civ:
                insts = page.search_for(marqueur_civ)
                for inst in insts:
                    page.draw_rect(inst, color=(1, 1, 1), fill=(1, 1, 1))
                    label = marqueur_civ.replace("q ", "[X] ")
                    page.insert_text(
                        (inst.x0, inst.y1 - 1),
                        label,
                        fontsize=8, color=(0, 0, 0), fontname="helv"
                    )
                    break

    # ── 6. Date du jour automatique ───────────────────────────────────────────
    from datetime import date as _date
    date_aujourd_hui = _date.today().strftime("%d/%m/%Y")
    for page_num in range(len(doc)):
        page = doc[page_num]
        insts = page.search_for("Date:")
        for inst in insts:
            page.insert_text(
                (inst.x1 + 5, inst.y1 - 1),
                date_aujourd_hui,
                fontsize=9, color=(0, 0, 0), fontname="helv"
            )
            break

    # ── 7. Sauvegarder ───────────────────────────────────────────────────────
    output = io.BytesIO()
    doc.save(output)
    doc.close()
    return output.getvalue()


def _remplir_docx_ameliore(docx_bytes, nom_fich, profil_enrichi, justifs):
    """
    Remplit un DOCX en séparant strictement labels et zones valeur.
    Claude retourne UNIQUEMENT les valeurs à insérer dans les zones vides,
    sans jamais écraser les labels existants.
    """
    doc = DocxDoc(io.BytesIO(docx_bytes))

    # Construire l'index avec type col (label=0 / valeur=1+)
    items = []

    for para in doc.paragraphs:
        if para.text.strip():
            items.append({
                "index": len(items),
                "type": "para",
                "texte": para.text,
                "ref": para,
                "col": -1
            })

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    if para.text.strip():
                        items.append({
                            "index": len(items),
                            "type": "cell",
                            "texte": para.text,
                            "ref": para,
                            "col": c_idx,
                            "row": r_idx,
                            "table": t_idx
                        })

    items_str = "\n".join(
        f"[{it['index']}] {'LABEL' if it['type']=='cell' and it['col']==0 else 'VALEUR' if it['type']=='cell' and it['col']>0 else 'PARA'} | {it['texte']}"
        for it in items
    )[:5000]

    justif_bloc = _construire_justif_bloc(justifs)

    system = (
        "Tu es un assistant expert en administration française. "
        "Tu analyses des documents administratifs et remplis les zones vides. "
        "Tu retournes UNIQUEMENT un objet JSON valide, "
        "sans texte avant ni après, sans balises markdown."
    )

    user = f"""Document administratif (LABEL=colonne gauche tableau à NE PAS modifier, VALEUR=colonne droite à remplir, PARA=paragraphe normal) :

{items_str}

Profil :
{json.dumps(profil_enrichi, ensure_ascii=False, indent=2)}
{justif_bloc}

Retourne un JSON :
{{
  "remplissages": [
    {{"index": 5, "valeur": "ALVAREZ"}},
    {{"index": 8, "valeur": "Clément"}}
  ]
}}

RÈGLES ABSOLUES :
1. N'inclure QUE les index "VALEUR" ou "PARA" contenant des pointillés (......) ou zones vides
2. JAMAIS inclure les index "LABEL" — ce sont les labels fixes du tableau
3. "valeur" = UNIQUEMENT la donnée, PAS le texte complet :
   - PARA "Nom :............." → valeur = profil["nom"] = NOM DE FAMILLE uniquement
   - PARA "Prénom :............." → valeur = profil["prenom"] = PRÉNOM uniquement
   - "(jj/mm/aa)" sur ligne Départ → valeur = profil["date_depart"] format jj/mm/aa
   - "(jj/mm/aa)" sur ligne Arrivée → valeur = profil["date_arrivee"] format jj/mm/aa

MAPPING TABLEAU VOYAGE — CRITIQUE :
- VALEUR sur la même ligne que LABEL "Train" → profil["frais_train"] (JAMAIS dans Autre)
- VALEUR sur la même ligne que LABEL "Avion" → profil["frais_avion"] si dispo
- VALEUR sur la même ligne que LABEL "Taxi" → profil["frais_taxi"] si dispo
- VALEUR sur la même ligne que LABEL "Péage" → profil["frais_peage"] si dispo
- VALEUR sur la même ligne que LABEL "Voiture personnel" → profil["frais_voiture"] si dispo
- VALEUR sur la même ligne que LABEL "Autre" → laisser vide sauf si frais "autre" existent
- VALEUR sur la même ligne que LABEL "Montant total" → profil["montant_total"]

4. Pour les cellules VALEUR du tableau : insérer uniquement le montant numérique (ex: "112.00")
5. Ne pas inventer d'information absente du profil
6. Si une info manque → ne pas inclure cet index

JSON :"""

    response = claude.messages.create(
        model=CLAUDE_MODEL, max_tokens=2000,
        system=system,
        messages=[{"role": "user", "content": user}]
    )

    raw = response.content[0].text.strip()
    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1] if len(parts) > 1 else raw
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip()

    try:
        result = json.loads(raw)
        remplissages = {str(r["index"]): r["valeur"] for r in result.get("remplissages", [])}
    except (json.JSONDecodeError, KeyError):
        print(f"WARNING: JSON invalide de Claude: {raw[:200]}")
        remplissages = {}

    PATTERNS_VIDE = [r'\.{5,}', r'…{3,}', r'_{5,}', r'\(jj/mm/(?:aa|aaaa)\)']

    for idx_str, valeur in remplissages.items():
        try:
            idx = int(idx_str)
            if idx >= len(items):
                continue
            item = items[idx]

            # Sécurité : jamais modifier un LABEL (col 0 d'un tableau)
            if item["col"] == 0:
                print(f"WARNING: tentative de modifier LABEL index={idx}, ignoré")
                continue

            para = item["ref"]
            _fusionner_runs(para)

            texte_actuel = para.text
            nouveau_texte = texte_actuel

            # Remplacer UNIQUEMENT les zones vides, pas tout le texte
            for pattern in PATTERNS_VIDE:
                if re.search(pattern, texte_actuel):
                    nouveau_texte = re.sub(pattern, str(valeur), texte_actuel, count=1)
                    break
            else:
                # Cellule vraiment vide → insérer la valeur
                if not texte_actuel.strip() or texte_actuel.strip() in ["", " "]:
                    nouveau_texte = str(valeur)

            if para.runs:
                para.runs[0].text = nouveau_texte
            else:
                from docx.oxml import OxmlElement
                r_el = OxmlElement('w:r')
                t_el = OxmlElement('w:t')
                t_el.text = nouveau_texte
                r_el.append(t_el)
                para._p.append(r_el)

        except Exception as e:
            print(f"WARNING: Erreur remplissage index {idx_str}: {e}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.route("/ai/remplir-document", methods=["POST"])
def ai_remplir_document():
    if not CLAUDE_OK:
        return jsonify({"error": "Service Claude non disponible"}), 503
    if not DOCX_OK:
        return jsonify({"error": "python-docx non installé"}), 503
    try:
        data     = request.get_json(force=True) or {}
        profil   = data.get("profil", {})
        b64      = data.get("fichier_base64", "")
        nom_fich = data.get("nom_fichier", "document.docx")
        justifs  = data.get("justificatifs", [])

        print(f"[REMPLIR] START nom={nom_fich} b64_len={len(b64)} nb_justifs={len(justifs)} fitz={FITZ_OK}")
        print(f"[REMPLIR] justifs_types={[j.get('type_document') for j in justifs]}")
        print(f"[REMPLIR] justifs_montants={[j.get('informations',{}).get('montant') for j in justifs]}")
        print(f"[REMPLIR] profil_keys={list(profil.keys())}")

        if not b64:
            return jsonify({"error": "fichier_base64 manquant"}), 400

        raw_bytes = base64.b64decode(b64)

        # ── Enrichir le profil ────────────────────────────────────────────────
        adresse_brute = profil.get("adresse", "")
        num_rue, nom_rue = _decomposer_adresse(adresse_brute)
        profil_enrichi = dict(profil)
        profil_enrichi["numero_rue"] = num_rue
        profil_enrichi["nom_rue"] = nom_rue
        profil_enrichi["adresse_complete"] = (
            f"{adresse_brute}, {profil.get('code_postal','')} "
            f"{profil.get('ville','')}".strip(", ")
        )

        # S'assurer que nom et prénom sont bien séparés (jamais fusionnés)
        # Si le profil a "nom_complet", le décomposer
        if profil_enrichi.get("nom_complet") and not profil_enrichi.get("nom"):
            parts = profil_enrichi["nom_complet"].split(" ", 1)
            profil_enrichi["prenom"] = parts[0] if len(parts) > 0 else ""
            profil_enrichi["nom"] = parts[1] if len(parts) > 1 else ""

        def _normaliser_date(d: str) -> str:
            """Convertit JJ/MM/AAAA vers jj/mm/aa."""
            if not d:
                return d
            match = re.match(r'(\d{1,2})/(\d{2})/(\d{4})', d)
            if match:
                return f"{match.group(1).zfill(2)}/{match.group(2)}/{match.group(3)[2:]}"
            return d

        # Enrichir avec justificatifs — calculer tous les types de frais
        totaux_frais = {}
        toutes_dates = []

        for j in justifs:
            ttype = j.get("type_document", "autre")
            info = j.get("informations", {})
            montant = info.get("montant")
            date = info.get("date", "")

            if montant is not None:
                try:
                    m_float = float(montant)
                    totaux_frais[ttype] = totaux_frais.get(ttype, 0.0) + m_float
                except (ValueError, TypeError):
                    pass

            if date and ttype not in ("rib",):
                toutes_dates.append(_normaliser_date(date))

        # Injecter dans profil_enrichi
        if totaux_frais.get("billet_train", 0) > 0:
            profil_enrichi["frais_train"] = f"{totaux_frais['billet_train']:.2f}"
        if totaux_frais.get("billet_avion", 0) > 0:
            profil_enrichi["frais_avion"] = f"{totaux_frais['billet_avion']:.2f}"
        if totaux_frais.get("taxi", 0) > 0:
            profil_enrichi["frais_taxi"] = f"{totaux_frais['taxi']:.2f}"
        if totaux_frais.get("peage", 0) > 0:
            profil_enrichi["frais_peage"] = f"{totaux_frais['peage']:.2f}"

        # Voiture : depuis le profil (km * 0.5) ou justificatif
        if profil.get("km_voiture"):
            try:
                km = float(profil["km_voiture"])
                profil_enrichi["frais_voiture"] = f"{km * 0.5:.2f}"
                profil_enrichi["km_voiture"] = str(int(km))
            except (ValueError, TypeError):
                pass

        # Montant total = somme de tous les frais sauf RIB
        total_global = sum(
            v for k, v in totaux_frais.items() if k != "rib"
        )
        # Ajouter la voiture si présente
        if profil.get("km_voiture"):
            try:
                total_global += float(profil["km_voiture"]) * 0.5
            except (ValueError, TypeError):
                pass
        if total_global > 0:
            profil_enrichi["montant_total"] = f"{total_global:.2f}"

        # Dates depuis justificatifs
        if toutes_dates:
            profil_enrichi["date_depart"] = toutes_dates[0]
            profil_enrichi["date_arrivee"] = toutes_dates[-1] if len(toutes_dates) > 1 else toutes_dates[0]

        # RIB
        rib = next((j for j in justifs if j.get("type_document") == "rib"), None)
        if rib:
            info_rib = rib.get("informations", {})
            if info_rib.get("iban"):
                profil_enrichi["iban"] = info_rib["iban"]
            if info_rib.get("bic"):
                profil_enrichi["bic"] = info_rib["bic"]
            if info_rib.get("titulaire"):
                profil_enrichi["titulaire_compte"] = info_rib["titulaire"]
            profil_enrichi["banque_nom"] = info_rib.get("banque", "La Banque Postale")

        ext = os.path.splitext(nom_fich)[1].lower()

        # ── LOG DIAGNOSTIC ────────────────────────────────────────────────────
        print(f"DEBUG profil_enrichi keys: {list(profil_enrichi.keys())}")
        print(f"DEBUG frais_train: {profil_enrichi.get('frais_train')}")
        print(f"DEBUG montant_total: {profil_enrichi.get('montant_total')}")
        print(f"DEBUG date_depart: {profil_enrichi.get('date_depart')}")
        print(f"DEBUG nb justifs reçus: {len(justifs)}")
        print(f"DEBUG justifs types: {[j.get('type_document') for j in justifs]}")
        print(f"DEBUG justifs montants: {[j.get('informations',{}).get('montant') for j in justifs]}")
        print(f"DEBUG totaux_frais: {totaux_frais}")
        # ─────────────────────────────────────────────────────────────────────
        if ext == ".pdf":
            if FITZ_OK:
                # Flux direct PyMuPDF — pas de conversion DOCX
                pdf_final = _remplir_pdf_direct(raw_bytes, nom_fich, profil_enrichi, justifs)
            else:
                # Fallback : conversion LibreOffice → DOCX → remplissage → PDF
                print("WARNING: PyMuPDF absent, fallback LibreOffice")
                try:
                    docx_bytes = _convert_pdf_to_docx(raw_bytes, nom_fich)
                    nom_fich = os.path.splitext(nom_fich)[0] + ".docx"
                except Exception as conv_err:
                    return jsonify({"error": f"Conversion PDF impossible: {conv_err}"}), 422
                docx_rempli = _remplir_docx_ameliore(docx_bytes, nom_fich, profil_enrichi, justifs)
                try:
                    pdf_final = _convert_docx_to_pdf(docx_rempli, nom_fich)
                except Exception as conv_err:
                    return jsonify({"error": f"Erreur conversion PDF final: {conv_err}"}), 500

        # ── BRANCHE DOC ───────────────────────────────────────────────────────
        elif ext == ".doc":
            try:
                docx_bytes = _convert_doc_to_docx(raw_bytes, nom_fich)
                nom_fich = os.path.splitext(nom_fich)[0] + ".docx"
            except Exception as conv_err:
                return jsonify({"error": "Conversion .doc impossible. Ouvrez le fichier dans Word et enregistrez-le en .docx."}), 422
            docx_rempli = _remplir_docx_ameliore(docx_bytes, nom_fich, profil_enrichi, justifs)
            try:
                pdf_final = _convert_docx_to_pdf(docx_rempli, nom_fich)
            except Exception as conv_err:
                return jsonify({"error": f"Erreur conversion PDF final: {conv_err}"}), 500

        # ── BRANCHE DOCX ──────────────────────────────────────────────────────
        else:
            docx_rempli = _remplir_docx_ameliore(raw_bytes, nom_fich, profil_enrichi, justifs)
            try:
                pdf_final = _convert_docx_to_pdf(docx_rempli, nom_fich)
            except Exception as conv_err:
                return jsonify({"error": f"Erreur conversion PDF final: {conv_err}"}), 500

        # ── Retourner le PDF ──────────────────────────────────────────────────
        p = profil_enrichi
        if p.get("nom") and p.get("prenom"):
            nom_sortie = f"document-rempli-{p['prenom'].lower()}-{p['nom'].lower()}.pdf"
        else:
            nom_sortie = f"{os.path.splitext(nom_fich)[0]}-complété.pdf"

        return send_file(
            io.BytesIO(pdf_final),
            mimetype="application/pdf",
            as_attachment=True,
            download_name=nom_sortie,
        )

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ══════════════════════════════════════════════════════════════════════════════
# IA — ANALYSER UN DOCUMENT ET IDENTIFIER LES CHAMPS MANQUANTS
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/ai/analyser-document", methods=["POST"])
def ai_analyser_document():
    if not CLAUDE_OK:
        return jsonify({"error": "Service Claude non disponible"}), 503
    raw = ""
    try:
        d      = request.get_json(force=True) or {}
        texte  = d.get("texte_document", "")
        profil = d.get("profil_existant", {})
        justifs = d.get("justificatifs", [])

        if not texte:
            return jsonify({"error": "texte_document manquant"}), 400

        # ── Extraire toutes les infos des justificatifs ───────────────────────
        infos_justifs = {}

        # Billets train — additionner tous les montants
        billets_train_analyse = [j for j in justifs if j.get("type_document") == "billet_train"]
        if billets_train_analyse:
            total_train = sum(
                float(j.get("informations", {}).get("montant") or 0)
                for j in billets_train_analyse
            )
            infos_justifs["frais_train"] = f"{total_train:.2f}"
            infos_justifs["montant_train"] = f"{total_train:.2f}"
            dates = [
                j.get("informations", {}).get("date")
                for j in billets_train_analyse
                if j.get("informations", {}).get("date")
            ]
            if dates:
                infos_justifs["date_depart"] = dates[0]
                infos_justifs["date_arrivee"] = dates[-1] if len(dates) > 1 else dates[0]

        # Billets avion
        billets_avion = [j for j in justifs if j.get("type_document") == "billet_avion"]
        if billets_avion:
            total_avion = sum(float(j.get("informations", {}).get("montant") or 0) for j in billets_avion)
            infos_justifs["frais_avion"] = f"{total_avion:.2f}"
            infos_justifs["montant_avion"] = f"{total_avion:.2f}"

        # Taxi
        taxis = [j for j in justifs if j.get("type_document") == "taxi"]
        if taxis:
            total_taxi = sum(float(j.get("informations", {}).get("montant") or 0) for j in taxis)
            infos_justifs["frais_taxi"] = f"{total_taxi:.2f}"
            infos_justifs["montant_taxi"] = f"{total_taxi:.2f}"

        # Péage
        peages = [j for j in justifs if j.get("type_document") == "peage"]
        if peages:
            total_peage = sum(float(j.get("informations", {}).get("montant") or 0) for j in peages)
            infos_justifs["frais_peage"] = f"{total_peage:.2f}"
            infos_justifs["montant_peage"] = f"{total_peage:.2f}"

        # Total global de tous les justificatifs
        total_global = sum(
            float(j.get("informations", {}).get("montant") or 0)
            for j in justifs
            if j.get("type_document") not in ("rib",)
        )
        if total_global > 0:
            infos_justifs["montant_total"] = f"{total_global:.2f}"

        # RIB
        for j in justifs:
            info = j.get("informations", {})
            if j.get("type_document") == "rib":
                if info.get("iban"):
                    infos_justifs["iban"] = info["iban"]
                if info.get("bic"):
                    infos_justifs["bic"] = info["bic"]
                if info.get("titulaire"):
                    infos_justifs["titulaire_compte"] = info["titulaire"]
                infos_justifs["banque_nom"] = info.get("banque", "La Banque Postale")

        # Fusionner profil + justificatifs — les justificatifs complètent le profil
        profil_complet = {**profil, **infos_justifs}
        champs_renseignes = {
            k: v for k, v in profil_complet.items()
            if v and str(v).strip() not in ("", "null", "None", "0", "0.00")
        }

        # Champs couverts par les justificatifs (ne pas redemander à l'utilisateur)
        CHAMPS_COUVERTS_PAR_JUSTIFS = set()
        if billets_train_analyse:
            CHAMPS_COUVERTS_PAR_JUSTIFS.update([
                "frais_train", "montant_train", "date_depart", "date_arrivee",
                "montant_total", "date_voyage"
            ])
        if billets_avion:
            CHAMPS_COUVERTS_PAR_JUSTIFS.update(["frais_avion", "montant_avion"])
        if taxis:
            CHAMPS_COUVERTS_PAR_JUSTIFS.update(["frais_taxi", "montant_taxi"])
        if peages:
            CHAMPS_COUVERTS_PAR_JUSTIFS.update(["frais_peage", "montant_peage"])
        if any(j.get("type_document") == "rib" for j in justifs):
            CHAMPS_COUVERTS_PAR_JUSTIFS.update(["iban", "bic", "titulaire_compte"])

        system = (
            "Tu es un assistant administratif expert. "
            "Tu analyses des documents administratifs français et identifies "
            "les champs personnels qui doivent être remplis par l'utilisateur. "
            "Tu retournes UNIQUEMENT un objet JSON valide, "
            "sans texte avant ni après, sans balises markdown."
        )

        champs_couverts_str = (
            f"\nChamps déjà couverts par les justificatifs (NE PAS inclure dans champs_manquants) : "
            f"{', '.join(sorted(CHAMPS_COUVERTS_PAR_JUSTIFS))}"
            if CHAMPS_COUVERTS_PAR_JUSTIFS else ""
        )

        user = f"""Voici le texte d'un document administratif à remplir :

{texte[:4000]}

Informations déjà disponibles (profil + justificatifs) :
{json.dumps(champs_renseignes, ensure_ascii=False, indent=2)}
{champs_couverts_str}

Identifie UNIQUEMENT les champs personnels présents dans ce document qui manquent encore.
Retourne ce JSON :
{{
  "champs_manquants": [
    {{
      "id": "identifiant_snake_case",
      "label": "Libellé clair pour l'utilisateur",
      "type": "texte" ou "choix" ou "date",
      "exemple": "exemple de valeur (optionnel)",
      "options": ["opt1", "opt2"]
    }}
  ],
  "champs_trouves": ["liste des champs déjà disponibles"]
}}

RÈGLES :
- Ne liste dans champs_manquants QUE les champs présents dans le document ET absents du profil
- Ne jamais demander les montants de transport si des justificatifs sont fournis
- Ne jamais demander les dates de voyage si des billets sont fournis
- Ne jamais demander l'IBAN/BIC si un RIB est fourni
- Pour "situation_famille" → type "choix", options : ["Célibataire", "Marié(e)", "Pacsé(e)", "Divorcé(e)", "Veuf/Veuve", "Concubinage"]
- Pour "situation_pro" → type "choix", options : ["Secteur public", "Secteur privé salarié", "Travailleur non-salarié / Auto-entrepreneur", "Intermittent du spectacle", "Étudiant(e) 3ème cycle", "Retraité(e)"]
- Pour "civilite" → type "choix", options : ["Pr", "Dr", "M.", "Mme", "Mlle"]

JSON :"""

        raw = claude_text(system, user, max_tokens=1500)
        raw = raw.strip()
        if raw.startswith("```"):
            parts = raw.split("```")
            raw = parts[1] if len(parts) > 1 else raw
            if raw.startswith("json"):
                raw = raw[4:]
        raw = raw.strip()

        result = json.loads(raw)
        if "champs_manquants" not in result:
            result["champs_manquants"] = []
        if "champs_trouves" not in result:
            result["champs_trouves"] = list(champs_renseignes.keys())
        return jsonify(result)

    except json.JSONDecodeError:
        print(f"WARNING: JSON invalide de Claude: {raw[:200]}")
        return jsonify({
            "champs_manquants": [
                {"id": "num_secu", "label": "N° de sécurité sociale", "type": "texte", "exemple": "1 85 03 75 056 789 42"},
                {"id": "iban", "label": "IBAN", "type": "texte", "exemple": "FR76 1234 5678 9012 3456 7890 123"},
                {"id": "situation_famille", "label": "Situation familiale", "type": "choix",
                 "options": ["Célibataire", "Marié(e)", "Pacsé(e)", "Divorcé(e)", "Veuf/Veuve", "Concubinage"]},
            ],
            "champs_trouves": list(profil.keys())
        })
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ══════════════════════════════════════════════════════════════════════════════
# PDF DEVIS
# ══════════════════════════════════════════════════════════════════════════════

try:
    from generate_devis import generate_devis_pdf
    DEVIS_OK = True
except ImportError as e:
    DEVIS_OK = False
    print(f"WARNING: generate_devis non trouvé: {e}")


@app.route("/generate-devis", methods=["POST"])
def generate_devis():
    if not DEVIS_OK:
        return jsonify({"error": "Module generate_devis non disponible"}), 503
    try:
        data   = request.get_json(force=True) or {}
        pdf    = generate_devis_pdf(data)
        dv     = data.get("devis", {})
        cl     = data.get("client", {})
        numero = dv.get("numero", "devis").replace(" ", "-").replace("/", "-")
        client = cl.get("nom", "client")[:20].replace(" ", "-")
        return send_file(io.BytesIO(pdf), mimetype="application/pdf",
                         as_attachment=True, download_name=f"Devis-{numero}-{client}.pdf")
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ══════════════════════════════════════════════════════════════════════════════
# PDF ORDRE DE MISSION
# ══════════════════════════════════════════════════════════════════════════════

try:
    from generate_ordre_mission import generate_ordre_mission_pdf
    MISSION_OK = True
except ImportError as e:
    MISSION_OK = False
    print(f"WARNING: generate_ordre_mission non trouvé: {e}")


@app.route("/generate-ordre-mission", methods=["POST"])
def generate_ordre_mission():
    if not MISSION_OK:
        return jsonify({"error": "Module generate_ordre_mission non disponible"}), 503
    try:
        data = request.get_json(force=True) or {}
        pdf  = generate_ordre_mission_pdf(data)
        mis  = data.get("missionnaire", {})
        nom  = mis.get("nom", "mission").replace(" ", "-").lower()
        msn  = data.get("mission", {})
        date = str(msn.get("date_debut", ""))[:10].replace("-", "")
        return send_file(io.BytesIO(pdf), mimetype="application/pdf",
                         as_attachment=True, download_name=f"OrdeMission-{nom}-{date}.pdf")
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ══════════════════════════════════════════════════════════════════════════════
# DÉMARRAGE
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    print(f"Démarrage sur port {port}")
    print(f"  Claude:{CLAUDE_OK} | Annexes:{ANNEXES_OK} | Vacataire:{VACATAIRE_OK} | Docx:{DOCX_OK} | Fitz:{FITZ_OK}")
    app.run(host="0.0.0.0", port=port, debug=False)
